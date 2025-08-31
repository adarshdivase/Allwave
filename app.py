import streamlit as st
import pandas as pd
import math
import re
import eml_parser
import docx
import io

# --- 1. Load Data & Core Functions ---
st.set_page_config(layout="wide")
st.title("🚀 All Wave AI-Powered Design & Estimation Engine")

@st.cache_data
def load_data():
    """Loads and prepares all necessary data files from the user's directory."""
    try:
        oem_df = pd.read_csv("av_oem_list_2025.csv")
        # Combine open and closed tickets for a comprehensive support database
        closed_tickets_df = pd.read_csv("Closed tickets(Last 10 days).xlsx - Jira Export Excel CSV (my defau.csv")
        open_tickets_df = pd.read_csv("Open Tickets(last 10 days).xlsx - Jira Export Excel CSV (my defau.csv")
        all_tickets_df = pd.concat([closed_tickets_df, open_tickets_df], ignore_index=True)
        all_tickets_df.rename(columns={'Summary': 'summary', 'Custom field (RCA - Root Cause Analysis)': 'rca'}, inplace=True)
        return oem_df, all_tickets_df
    except FileNotFoundError as e:
        st.error(f"Fatal Error: {e}. Please ensure all required CSV files are in the same folder as this script.")
        return None, None

oem_list_df, tickets_df = load_data()

def parse_docx(file_stream):
    """AI-powered parser for .docx Need Analysis Documents."""
    try:
        doc = docx.Document(file_stream)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Regex to find key values like "Length - 20ft" or "28 Pax"
        length_ft = re.search(r'Length\s*-\s*(\d+(\.\d+)?)ft', text, re.IGNORECASE)
        capacity = re.search(r'(\d+)\s*Pax', text, re.IGNORECASE)
        
        return {
            "room_name": doc.paragraphs[0].text.strip().replace(":", ""),
            "farthest_viewer": float(length_ft.group(1)) * 0.3048 if length_ft else 6.0, # Convert ft to m
            "capacity": int(capacity.group(1)) if capacity else 12,
        }
    except Exception:
        st.warning("Could not fully parse the DOCX. Please review the extracted details.")
        return None

# --- 2. The AI Proposal Generation Engine ---
def generate_proposal(capacity, farthest_viewer, use_case, has_direct_light, tier="Standard"):
    """The core engine that runs compliance checks and generates tiered BOQs."""
    # --- A. AVIXA Compliance Check ---
    # DISCAS Standard for Display Size
    min_height_m = farthest_viewer / 15 if use_case == "Analytical Decision Making" else farthest_viewer / 20
    min_diagonal_inches = min_height_m * 39.37 * 1.89
    sizes = [55, 65, 75, 85, 98]
    rec_size = min(sizes, key=lambda x:abs(x-min_diagonal_inches))
    discas_report = f"Recommended display is {rec_size}\" based on viewing distance."

    # Lighting & Glare Check
    light_report, light_status = ("FAIL: High risk of screen glare.", "❌") if has_direct_light else ("PASS: No lighting conflicts.", "✅")

    compliance_df = pd.DataFrame([
        {"Standard": "Display Size (DISCAS)", "Result": discas_report, "Status": "✅"},
        {"Standard": "Lighting & Glare", "Result": light_report, "Status": light_status}
    ])

    # --- B. Tiered BOQ Generation ---
    brand_tiers = {
        "Budget": ["Yealink", "BenQ", "Aver"],
        "Standard": ["Logitech", "Poly", "Samsung"],
        "Premium": ["Cisco", "Crestron", "Shure", "Biamp"]
    }
    
    if capacity <= 6: template = ["UC & Collaboration Devices", "Displays & Projectors"]
    elif capacity <= 14: template = ["UC & Collaboration Devices", "Displays & Projectors", "Audio: Microphones & Conferencing"]
    else: template = ["PTZ & Pro Video Cameras", "Video Wall Technology", "Audio: DSP & Mixing", "Control Systems", "Acoustics & Sound Masking"]

    boq_items = []
    for category in template:
        brands = oem_list_df[oem_list_df['Category'] == category]
        selected_brand = brands[brands['OEM / Brand'].str.contains('|'.join(brand_tiers[tier]), case=False)]
        brand = selected_brand['OEM / Brand'].iloc[0] if not selected_brand.empty else (brands['OEM / Brand'].iloc[0] if not brands.empty else "N/A")
        desc = f"{rec_size}\" Display" if "Display" in category else f"{brand} Solution"
        boq_items.append({"Category": category, "Brand Recommendation": brand, "Description": desc, "Tier": tier, "Qty": 1})

    return compliance_df, pd.DataFrame(boq_items)

# --- 3. Streamlit User Interface ---
if oem_list_df is not None:
    # Initialize session state for variables
    for key in ['room_name', 'capacity', 'farthest_viewer']:
        if key not in st.session_state: st.session_state[key] = ""

    st.sidebar.header("1. Start a New Project")
    st.sidebar.info("Upload a Need Analysis Doc (.docx) to auto-fill.")
    
    uploaded_file = st.sidebar.file_uploader("Upload Document", type=['docx'])
    
    if uploaded_file:
        extracted = parse_docx(io.BytesIO(uploaded_file.getvalue()))
        if extracted:
            st.session_state.update(extracted)
            st.sidebar.success(f"Analyzed: **{st.session_state.room_name}**")

    # --- Input Form in Sidebar ---
    st.sidebar.text_input("Room Name", key="room_name")
    st.sidebar.number_input("Seating Capacity", min_value=1, key="capacity")
    st.sidebar.number_input("Farthest Viewer (meters)", key="farthest_viewer")
    use_case = st.sidebar.selectbox("Primary Use Case", ["Analytical Decision Making", "Basic Decision Making"])
    has_light = st.sidebar.checkbox("Light source above display?")

    if st.sidebar.button("🚀 Generate AI Proposal"):
        st.session_state.compliance, st.session_state.boq = generate_proposal(st.session_state.capacity, st.session_state.farthest_viewer, use_case, has_light, "Standard") # Default to standard
        st.session_state.proposal_generated = True

    # --- Main Display Area ---
    st.header("2. AI-Generated Proposal")
    if 'proposal_generated' in st.session_state:
        # Tier selection buttons
        c1, c2, c3 = st.columns(3)
        if c1.button("💵 Generate Budget BOQ"):
            _, st.session_state.boq = generate_proposal(st.session_state.capacity, st.session_state.farthest_viewer, use_case, has_light, "Budget")
        if c2.button("⭐ Generate Standard BOQ"):
            _, st.session_state.boq = generate_proposal(st.session_state.capacity, st.session_state.farthest_viewer, use_case, has_light, "Standard")
        if c3.button("💎 Generate Premium BOQ"):
            _, st.session_state.boq = generate_proposal(st.session_state.capacity, st.session_state.farthest_viewer, use_case, has_light, "Premium")
        
        st.subheader("Bill of Quantities (BOQ)")
        st.table(st.session_state.boq)
        
        st.subheader("AVIXA Compliance Report")
        st.dataframe(st.session_state.compliance.style.apply(lambda row: ['color:red' if row.Status == '❌' else '' for v in row], axis=1))
    else:
        st.info("Upload a document or fill in the details on the left and click 'Generate'.")

    # --- Support Chatbot ---
    st.header("3. Historical Support Ticket Search")
    query = st.text_input("Search past tickets (e.g., 'projector image', 'Crestron'):")
    if query and tickets_df is not None:
        results = tickets_df[tickets_df.apply(lambda r: query.lower() in str(r['summary']).lower() or query.lower() in str(r['rca']).lower(), axis=1)]
        if not results.empty:
            st.dataframe(results[['Issue key', 'Status', 'summary', 'rca']].head())
